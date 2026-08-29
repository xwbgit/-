import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from pathlib import Path

def create_competition_report_docx():
    doc = docx.Document()
    
    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 封面
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_top = p_top.add_run("附件5\n第九届浙江省大学生网络与信息安全竞赛\n作品挑战赛作品报告")
    run_top.font.name = "黑体"
    run_top.font.size = Pt(12)
    run_top.font.color.rgb = RGBColor(100, 100, 100)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(80)
    p_title.paragraph_format.space_after = Pt(20)
    run_t = p_title.add_run("DAS-SentinelAgent：面向网站安全风险评估与敏感信息防泄露的智能巡检智能体系统")
    run_t.font.name = "黑体"
    run_t.font.size = Pt(22)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(11, 87, 208)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(100)
    run_sub = p_sub.add_run("【企业命题 2：杭州安恒信息技术股份有限公司】")
    run_sub.font.name = "宋体"
    run_sub.font.size = Pt(14)

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.line_spacing = 1.8
    r_info = p_info.add_run("参赛赛道：企业命题\n作品名称：DAS-SentinelAgent 网站安全智能巡检智能体\n提交日期：2026年10月")
    r_info.font.name = "宋体"
    r_info.font.size = Pt(12)

    doc.add_page_break()

    # 样式配置辅助函数
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        r = h.add_run(text)
        r.font.name = "黑体"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = "黑体"
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)
        return h

    def add_body_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "宋体"
            rb.font.size = Pt(12)
            rb.bold = True
        r = p.add_run(text)
        r.font.name = "宋体"
        r.font.size = Pt(12)
        return p

    # 摘要
    h_abs = add_heading_1("摘  要")
    h_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body_p("政企网站长期面临组件漏洞、配置缺陷、页面篡改挂马和敏感数据误发风险。本项目实现了 DAS-SentinelAgent 可部署原型，提供授权边界内的页面发现、非破坏性内置探针、暗链挂马检查、敏感信息校验、去重定级、基线对比、告警、报告和复测闭环；同时提供默认关闭的 Nuclei、Gitleaks 与 OWASP ZAP 适配层。在版本化本地标注集的 19 个正样本和 4 个负样本上，当前可复现结果为 TP=19、FP=0、FN=0、TN=4。该结果只对内置小样本靶场有效，不代表公网通用效果；AI 动态规划和恒脑平台运行验证不在当前功能修复范围内。")

    doc.add_page_break()

    # 第一章 作品概述
    add_heading_1("第一章 作品概述")
    add_heading_2("1.1 研发背景与行业挑战")
    add_body_p("随着数字化改革的深入推进，政企门户网站及对外服务系统已成为公共服务与业务运转的核心窗口。然而，政企网站长期暴露于互联网复杂环境中，极易因 Web 组件漏洞、安全配置缺陷、页面被恶意篡改或植入暗链/挖矿脚本而遭受攻击；与此同时，工作人员因操作失误直接发布包含居民身份证、手机号、银行卡号或系统秘钥的通知公告，导致重大敏感信息泄露事件频发。")
    add_body_p("当前传统安全运维面临三大核心挑战：一是传统人工巡检周期长、人力成本高、覆盖面狭窄；二是传统扫描工具各自为战，缺乏协同编排能力且容易产生破坏性探测，影响业务连续性；三是缺乏持续运营闭环机制，无法量化基线变化与修复效果。")

    add_heading_2("1.2 项目目标与核心任务")
    add_body_p("本项目紧密围绕安恒信息企业赛题答题要求，构建具备自主规划、工具编排、智能研判与持续闭环运营能力的网站安全智能巡检智能体原型：")
    add_body_p("1. 授权范围与敏感规则自定义：支持录入授权域名、巡检深度/QPS与企业特定敏感信息定义；\n2. 自动发现与开源工具编排：自动递归发现站点页面与资源，编排安全探针矩阵识别漏洞、弱配置、暗链挂马及数据泄露；\n3. 智能研判与非破坏性证据链：智能去重、多级定级，提取上下文掩码证据链与代码级修复建议；\n4. 持续运营与基线对比：支持定时调度、历史基线 Diff 比对、多渠道告警与标准合规报告导出；\n5. 安恒恒脑生态无缝兼容：原生提供符合安恒恒脑安全智能体平台标准的 Tool Manifest 接口。")

    # 第二章 作品设计与实现
    add_heading_1("第二章 作品设计与实现")
    add_heading_2("2.1 系统总体架构设计")
    add_body_p("DAS-SentinelAgent 采用分层解耦的云原生 Agent 架构，包含接入层、智能体规划中枢 (Brain & Orchestrator)、多维探针检测矩阵 (Probes Matrix)、持续安全运营闭环引擎以及数据持久与审计层。")
    
    add_heading_2("2.2 核心功能模块设计与实现")
    add_body_p("（1）资产与页面深度发现引擎 (AssetCrawler)：采用异步并发 BFS 拓扑遍历，提取 HTML、JS 动态接口与静态资源，并在发起任何探测前强制执行 is_authorized() 授权边界过滤，严禁非授权越界探测。")
    add_body_p("（2）漏洞与弱配置非破坏性探针 (VulnDetector)：覆盖 OWASP 安全响应头缺失 (HSTS/CSP/X-Frame)、CORS 任意 Origin 反射缺陷、.git/HEAD、.env、backup.sql、Swagger、Actuator 等关键资产暴露，采用非破坏性轻量探针验证。")
    add_body_p("（3）页面防篡改与挂马暗链检测引擎 (TamperDetector)：深度解析 DOM 树，精准识别 display:none、负坐标绝对定位等隐蔽暗链，识别涉黑博彩外链、Hacked by 涂鸦篡改以及 coinhive/eval 混淆挂马脚本。")
    add_body_p("（4）深度敏感信息与隐私防泄露引擎 (SensitiveInspector)：支持身份证 ISO 7064:1983.MOD 11-2 校验位算法、银行卡 Luhn 模 10 算法、11位手机号号段校验、云厂商 AK/SK 及数据库连接串，并支持企业自定义正则沙箱测试与自动脱敏展示。")
    add_body_p("（5）基线差异比对与闭环引擎 (BaselineService)：利用 DOM Hash 与漏洞指纹对比两次快照，精确呈现'新增隐患'、'已闭环修复隐患'与'页面异动'，驱动安全态势持续收敛。")

    add_heading_2("2.3 恒脑安全智能体对接规范")
    add_body_p("仓库保留了恒脑工具定义与相关接口，本轮功能修复未修改该接口，也未在恒脑平台进行账号、权限、调用链和运行时兼容性验收。正式报送前应以实际平台运行记录补充本节。")

    # 第三章 作品测试与分析
    add_heading_1("第三章 作品测试与分析")
    add_heading_2("3.1 统一测试集与测试环境")
    add_body_p("为了量化评估检测能力，团队构建了高度仿真的政企典型安全缺陷统一测试靶场 (http://127.0.0.1:8088)，内置 19 项涵盖高危漏洞、配置缺陷、隐藏暗链、挖矿脚本、涂鸦篡改及敏感数据（含合法与伪造干扰样本）的真实用例。")

    add_heading_2("3.2 评测指标与结果分析")
    add_body_p("在版本化本地标注集上运行端到端 Pytest，当前结果为 TP=19、FP=0、FN=0、TN=4，因此 Precision=1.0、Recall=1.0、F1=1.0、FPR=0.0。指标由 backend/app/evaluation/metrics.py 根据 tests/fixtures/local_lab_ground_truth.json 自动计算，仅适用于这 23 个本地已标注样本，不应外推为系统在互联网场景的通用准确率。")

    # 第四章 创新性和特色说明
    add_heading_1("第四章 创新性和特色说明")
    add_body_p("1. 校验位驱动的敏感数据降误报：对身份证、银行卡和手机号在正则命中后执行结构与校验规则复核。")
    add_body_p("2. 可审计的策略编排：记录每个内置探针和可选外部工具的执行、跳过或失败原因；AI 动态规划仍属后续工作。")
    add_body_p("3. 非破坏性与授权边界：通过授权域校验、速率、超时和日志审计降低对目标业务的影响，生产使用前仍需由授权单位完成容量与变更评估。")
    add_body_p("4. 持续安全运营基线闭环：基于快照 Diff 实现从发现、告警、整改到复测闭环的一站式管理。")

    # 第五章 竞品分析
    add_heading_1("第五章 竞品分析")
    add_body_p("本原型的差异化方向是将页面发现、暗链挂马、敏感信息校验、外部工具结果、基线和复测统一到同一任务记录。当前尚未完成与商业平台或主流扫描器的对照实验，不对准确率、性能或平台兼容性做未验证的优劣结论。")

    # 第六章 总结
    add_heading_1("第六章 总结")
    add_body_p("本项目已形成可部署的本地功能原型，完成授权、发现、检测、去重、证据、基线、告警、报告与复测的主要链路。在正式参赛提交前，仍需完成更大标注集评估、真实外部工具环境验收、AI 动态规划以及恒脑平台运行验证。")

    # 第七章 附件
    add_heading_1("第七章 附件")
    add_body_p("附件包括：系统源代码包、一键启动脚本 (start.bat)、Docker 镜像配置文件、OpenAPI 接口规范说明文档及统一测试集评估报告。")

    output_path = r"d:\Gemini Work\DAS_SentinelAgent\docs\第九届浙江省大学生网络与信息安全竞赛作品挑战赛作品报告_企业命题2.docx"
    doc.save(output_path)
    print("Report docx generated successfully at:", output_path)

if __name__ == "__main__":
    create_competition_report_docx()
